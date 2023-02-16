# utf-8


import os
import shutil
import zipfile
import docx

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from clear_non_read_symbol import clear_non_read_symbol


def add_path_folder():
    """ввести путь до папки с файлами, прочитать папку, в цикле взять файл, прочитать его по параграфам"""
    print('введите путь до папки с файлами')
    path_folder = 'C:\\Users\\User\\PycharmProjects\\minimod\\Upload_folder'  # input()
    # path_folder = '/home/mannur/PycharmProjects/minimod/Upload_folder'
    if path_folder:
        return path_folder


def input_data():
    """ввод данных для изменения в файле"""
    print('введите текст поиска')
    search_string = input()
    print('введите текст для замены')
    replace_string = input()
    return search_string, replace_string


def unzip_archive():
    extract_dir = 'Upload_folder'
    file_name = '_instr.zip'
    # with os.scandir(extract_dir) as scan:
    #     return next(scan, None) is None
    if len(os.listdir(extract_dir)) > 0:
        shutil.rmtree(extract_dir)
        os.mkdir(extract_dir)
    print('0')
    with zipfile.ZipFile(file_name) as zf:
        zf.extractall(extract_dir)
    return


def read_files(path_folder):
    """"переименование файлов в латиницу по одному"""
    folder = path_folder
    folder_base = os.getcwd()
    list_old = []
    for file_name in os.listdir(folder):
        count_name = ''
        count_simbol = 0
        for simbol in file_name:
            if count_simbol == 3:
                break
            if simbol.isdigit():
                count_name += simbol
                count_simbol += 1
        file_rename = 'instr' + '_' + count_name + '.docx'
        if file_rename not in os.listdir(folder):
            os.chdir(folder)
            os.rename(file_name, file_rename)
            os.chdir(folder_base)
        else:
            list_old.append(file_name)
    return


def compile_file(path_folder):
    def cell_format():
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(12)
        cell.paragraphs[0].runs[0].font.bold = True
        return

    for file_name in os.listdir(path_folder):
        print('6,5', file_name)
        symbol_s = '\\'  # '/'
        # symbol_s = '/'
        text_txt = ''
        if file_name:
            print('6,6', file_name)
            if file_name[-5:] == '.docx':
                print('6,7', file_name)
                base_file_name = file_name[:-5]
                file_txt = base_file_name + '.txt'
                with open(file_txt, encoding='utf-8') as f_txt:
                    text_txt = f_txt.readlines()
                    print('text_txt')
            else:
                print('6,75', file_name)
                continue
        else:
            print('6,8', file_name)
            continue
        # for txt in text_txt:
        #     txt = txt[:-1]

        print('6,9', file_name)
        path_file = path_folder + symbol_s + file_name
        doc = docx.Document(path_file)  # нужен??
        doc_new = docx.Document()
        style = doc_new.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        anchor = '0'
        for index, string in enumerate(text_txt):
            if anchor == '0':
                print('Table')
                table = doc_new.add_table(rows=4, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = table.cell(0, 0)  # получаем ячейку таблицы
                # записываем в ячейку данные
                cell.text = 'Государственное бюджетное общеобразовательное учреждение'
                cell_format()  # форматируем по образцу ячейку
                cell = table.cell(1, 0)
                cell.text = 'Самарской области средняя общеобразовательная школа'
                cell_format()
                cell = table.cell(2, 0)
                cell.text = '«Образовательный  центр имени В.Н. Татищева» с. Челно-Вершины'
                cell_format()
                cell = table.cell(3, 0)
                cell.text = 'муниципального района Челно-Вершинский Самарской области'
                cell_format()
                table_1 = doc_new.add_table(rows=3, cols=2)
                cell = table_1.cell(0, 0)
                cell.text = 'Согласовано'
                cell_format()
                cell = table_1.cell(0, 1)
                cell.text = 'Утверждаю'
                cell_format()
                cell = table_1.cell(1, 0)
                cell.text = 'Председатель профкома'
                cell_format()
                cell = table_1.cell(2, 0)
                cell.text = '______________Сергеева Н.А.'
                cell = table_1.cell(1, 1)
                cell.text = 'Директор школы'
                cell_format()
                cell = table_1.cell(2, 1)
                cell.text = '______________Моисеева Н.В.'
                cell_format()

                table_2 = doc_new.add_table(rows=1, cols=2)
                cell = table_2.cell(0, 1)
                cell.text = 'Приказ № 142-од от 01.03.2022'
                cell_format()
                doc_new.add_paragraph()
                para = doc_new.add_paragraph()
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Центрирование заголовков по центру`
                para.add_run(string).bold = True
                print('7')
                anchor = '1'

            else:
                p_text = string
                para = doc_new.add_paragraph()

                list_format_center = ["ИОТ", "Инс", "ИНС",
                                      'I. ', 'II.', 'III', 'IV.', 'V. ', 'VI.', 'VII', 'IX.', 'X. ', 'XI.', 'XII',
                                      'XV.', 'XVI']
                arab_symbol = ['1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ']
                arab_symbol_2 = ['10. ', '11. ', '12. ', '13. ', '14. ', '15. ', '16. ', '17. ', '18. ', '19. ']

                if p_text[:3] in list_format_center:
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Центрирование заголовков по центру`
                    p_text = string.replace("\r", "")
                    # p_text = string.replace("\n", "")
                    para.add_run(p_text).bold = True

                elif p_text[:3] in arab_symbol or p_text[:4] in arab_symbol_2:
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    p_text = string.replace("\r", "")
                    # p_text = string.replace("\n", "")
                    para.add_run(p_text).bold = True

                if p_text[:3] == "ИНС" or p_text[:3] == "Инс":
                    doc_new.add_paragraph(text_txt[index + 1])
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    p_text = string.replace("\r", "")
                    p_text = string.replace("\n", "")
                    # para.add_run(p_text).bold = True
                    anchor = '2'

                else:
                    if anchor == '2':
                        anchor = '1'
                        continue
                    p_text = string.replace("\r", "")
                    p_text = string.replace("\n", "")
                    para.text = p_text
                    para.alignment = 3  # выравниевание по ширине
                    run = para.add_run()
                    para.paragraph_format.line_spacing = 1.0
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

        doc_new.save('new_' + file_name)
        os.remove(file_txt)


def main():
    print('1')
    path_folder = add_path_folder()
    print('2')
    os.chdir(path_folder)
    print('3')
    read_files(path_folder)
    print('4')
    clear_non_read_symbol(path_folder)
    print('5')
    compile_file(path_folder)
    print('end')


if __name__ == '__main__':
    main()

# TODO Нужно расписать максимальнов возможные ошибки для выявления и коррекции
# Нужно сделать единый шаблон для инструкции и на основании его привести все  инструкции к единообразию!


#  Удалить лишние интервалы?
# def delete_paragraph(paragraph):  # Удалить пустой абзац
#     p = paragraph._element
#     p.getparent().remove(p)
#     p._p = p._element = None

# Порядок -  находим файл, считываем его, конвертируем в тхт, создаем новый файл, добавляем в него шапку,
# добавляем обзацы и текст в новом формате,  сохраняем в новой папке.

# if p_text[-1:] == '^|':
#     p_text[-1:] = '^p'
#     print('Замена мягкого абзаца')

#  docx2txt спользовать очищенный текст
