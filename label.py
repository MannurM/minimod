import os
import docx
from docx import Document

path_folder = 'C:\\Users\\User\\PycharmProjects\\minimod\\Upload_folder'  # input()
# path_folder = '/home/mannur/PycharmProjects/minimod/Upload_folder'

# symbol = '/'  #
symbol = '\\'
print('Yfxfkb')
for file_name in os.listdir(path_folder):
    print('1')
    if file_name[-5:] != '.docx':
        continue

    print(file_name)
    list_new = []
    doc = Document(path_folder + symbol + file_name)
    all_paragr = doc.paragraphs

    for par in all_paragr:
        run = par.add_run()
        text = par.text
        new_text = text
        del_label = ['\\r', '\\n', '\n', '\r']
        del_label_bef = [' ', '\\t', '\t']
        list_index = []
        for index, txt in enumerate(text):
            if text[index] in del_label:
                list_index.append(index)
            if text[index].isalpha() or text[index].isdigit():
                # print('exit')
                break
            if text[0] in del_label_bef:
                list_index.append(index)

        for i in list_index:
            print(f' simbol-{ord(text[i])}-', text)
        par.add_break = False

    doc.save(path_folder + symbol + file_name)

