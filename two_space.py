import os
import docx
import docx2txt
from docx import Document

path_folder = 'C:\\Users\\User\\PycharmProjects\\minimod\\Upload_folder'  # input()
# path_folder = '/home/mannur/PycharmProjects/minimod/Upload_folder'

# symbol = '/'  #
symbol = '\\'
print('Yfxfkb')
for file_name in os.listdir(path_folder):
    print('1')
    if file_name[-5:] != '.docx':
        continue

    print(file_name)
    list_new = []
    doc = Document(path_folder + symbol + file_name)
    all_paragr = doc.paragraphs

    for par in all_paragr:
        # run = par.add_run()
        text = par.text
        new_text = text
        if not text:
            continue
        list_index = []
        label = 0
        for index, txt in enumerate(text):
            if text[index].isalpha() or text[index].isdigit():
                # print('exit')
                break
            elif text[:1] == ' ':
                text = text.lstrip()
        text = text.rstrip()
        if text[-2:] == '\r':
            text = text.replace('\\r', '')
        if text[-1:] == '\r':
            text = text.replace('\r', '')

        print('LF')
        par.text = text
        count = len(text) + 1
        index = 0
        print('count before', count)
        while count > 1:
            print(count, text[index])
            if text[index] == ' ' and text[index + 1] == ' ':
                text = text.replace('  ', ' ')
                par.text = text
                print(text)
            count -= 1
            index += 1
    print('end')




        # for index, txt in enumerate(text + ' '):
        #     if text[index] and text[-1:]:
        #         if text[index] == ' ' and text[index + 1] == " ":
        #             text = text[index].replace(text[index], '')
        # par.text = text

    doc.save(path_folder + symbol + file_name)