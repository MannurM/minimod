import os
import docx
from docx import Document

path_folder = 'C:\\Users\\User\\PycharmProjects\\minimod\\Upload_folder'  # input()
# path_folder = '/home/mannur/PycharmProjects/minimod/Upload_folder'

# symbol = '/'  #
symbol = '\\'

list_new = []
file_name = 'instr_110.docx'


def delete_paragraph(paragraph):  # удаление пустого параграфа
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def analiz(file_name):
    doc = Document(path_folder + symbol + file_name)
    all_paragr = doc.paragraphs
    count_par, count_new = 0, 0
    print('1')
    for par in all_paragr:
        # run = par.add_run()
        text = par.text
        # style_par = par.style
        new_text = text
        del_label = ['\\r', '\\n', '\n', '\r']
        del_label_bef = [' ', '\\t', '\t']
        list_index = []
        count_par += 1
        print('par')
        for index, txt in enumerate(text):
            print('txt')
            if text[index:index + 3] == "ИНС" or text[index:index + 3] == "Инс":
                print('cont')
                count_new = count_par
    for par in all_paragr:
        text = par.text
        for index, txt in enumerate(text):
            if text[index:index + 3] == "ИНС" or text[index:index + 3] == "Инс":
                break
        if count_new <= 1:
            break
        else:
            print('3')
            par.clear()
            delete_paragraph(par)
            count_new -= 1
    print('4')
    doc.save(path_folder + symbol + 'new_' + file_name)


analiz(file_name)