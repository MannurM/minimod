# utf-8
# скрипт по созданию шаблона файлов инструкций docx



# TODO делать ли Web interface
# TODO FLASK?

# запрос значений полей и сохранение в файл
# распаковка  папки зип архив из своих файлов
#   чтение файла из папки
#   поиск мест вставки замены
#   вставка данных
#   сохранение
# выход зип архив с изменениями


import docx
import zipfile
import os
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Mm


def input_data():
    """ввод данных для изменения в файле"""
    print('введите текст поиска')
    d_input, d_replace = {}, {}

    d_input['name_org'] = """ Государственное бюджетное общеобразовательное учреждение\
                          Самарской области средняя общеобразовательная школа \
                          «Образовательный  центр имени В.Н. Татищева» с. Челно-Вершины \
                          муниципального района Челно-Вершинский Самарской области"""
    d_input['name_org_var1'] = 'Самарской области '
    d_input['name_org_var2'] = 'средняя общеобразовательная школа'
    d_input['name_org_var3'] = 'муниципального района'
    d_input['name_org_var4'] = 'Государственное бюджетное'
    d_input['name_org_var5'] = 'Образовательный  центр '

    d_input['name_prof'] = 'Сергеева'
    d_input['name_director'] = 'Моисеева'
    d_input['name_spec'] = 'Зайдуллин'
    d_input['name_pos_prof'] = "Председатель профкома"
    d_input['name_pos_direct'] = 'Директор школы'
    d_input['name_prof_var1'] = 'Н.А.Сергеева'
    d_input['name_director_var1'] = 'Н.В.Моисеева'
    d_input['grif_1'] = 'СОГЛАСОВАНО'
    d_input['grif_2'] = 'УТВЕРЖДАЮ'
    d_input['data'] = '25 января 2017 г.'

    print('введите текст изменения!')
    d_replace['name_org'] = ''
    d_replace['name_org_var1'] = ''
    d_replace['name_org_var3'] = ''
    d_replace['name_org_var4'] = ''
    d_replace['name_org_var2'] = ''
    d_replace['name_org_var5'] = ''
    d_replace['name_prof'] = ''
    d_replace['name_director'] = ''
    d_replace['name_pos_prof'] = ''
    d_replace['name_pos_direct'] = ''
    d_replace['name_prof_var1'] = ''
    d_replace['name_director_var1'] = ''
    d_replace['name_spec'] = ''
    d_replace['grif_1'] = ''
    d_replace['grif_2'] = ''
    d_replace['data'] = ''
    return d_input, d_replace


def del_paragrafs_docx():
    folder = 'Upload_folder'
    for file in os.listdir(folder):
        file_path = os.path.join(os.getcwd(), folder, file)
        doc = docx.Document(file_path)
        all_paragrs = doc.paragraphs
        all_tables = doc.tables
        count = 1
        for paragr in all_paragrs:
            if "№" in paragr.text:
                break
            else:
                count += 1
        for paragr in all_paragrs:
            delete_paragraph(paragr)
            if count <= 2:
                break
            else:
                count -= 1
        for table in all_tables:
            if table.cell(0, 0).paragraphs[0].text == '':
                table._element.getparent().remove(table._element)
                print('Table delete!')
        doc.save(file_path)


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None
    return


def general_format():
    pass
    """приведение текста к единому стилю всех документов"""
    folder = 'Upload_folder'
    set_styles = set()
    # for file in os.listdir(folder):
    #     file_path = os.path.join(os.getcwd(), folder, file)
    #     doc = docx.Document(file_path)
        # all_paragrs = doc.paragraphs
        # all_tables = doc.tables

        # general_index = ['1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ', '10. ', '11. ',  '12. ']
        # print(file)
        # all_styles = doc.styles
        # paragraph_styles = [s for s in all_styles if s.type == WD_STYLE_TYPE.PARAGRAPH]

        # for style in paragraph_styles:
        #     print('style.name', style.name)
        #     set_styles.add(style.name)
        # for paragr in all_paragrs:
        #     if 'Общие' in paragr.text:
        #         print('1')  # стиль 1
        #     # if general_index in paragr.text:
        #     #     print('12')
        #     if '№' in paragr.text:
        #         print('13')
        #     if 'ИОТ' in paragr.text:
        #         print('14')

    for file in os.listdir(folder):
        file_path = os.path.join(os.getcwd(), folder, file)
        doc = docx.Document(file_path)
        all_paragrs = doc.paragraphs
        all_tables = doc.tables
        if 'UserHead1' not in set_styles:
            style = doc.styles.add_style('UserHead1', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Times New Roman'
            style.font.size = Pt(14)
            print('Стиль Добавлен!')

        style = doc.styles['UserHead1']
        if style.font.underline:
            style.font.underline = False

        for paragr in all_paragrs:
            paragr.style = 'UserHead1'
            paragr.paragraph_format.space_before = Mm(2)
            paragr.paragraph_format.space_after = Mm(2)
        print(file)
        doc.save(file_path)


def unzip_archive():
    extract_dir = 'Upload_folder'
    file_name = '_instr_.zip'
    with zipfile.ZipFile(file_name) as zf:
        zf.extractall(extract_dir)
    return


def replace_data(d_input, d_replace):
    """поиск текста на замену в файле"""
    folder = 'Upload_folder'
    for file in os.listdir(folder):
        file_path = os.path.join(os.getcwd(), folder, file)
        doc = docx.Document(file_path)
        all_paragrs = doc.paragraphs
        for key, value in d_input.items():
            for paragr in all_paragrs:
                inline = paragr.runs
                if value in paragr.text:
                    print('Sucsess!')
                    paragr.text = d_replace[key]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if value in cell.text:
                            print('Sucsess! Table!')
                            cell.text = d_replace[key]
        doc.save(file_path)
    return


def zip_template():
    path = 'Upload_folder'  # '/home/docs-python/script/sql-script/'
    file_dir = os.listdir(path)
    with zipfile.ZipFile('_instr_template.zip', mode='w', compression=zipfile.ZIP_DEFLATED) as zf:
        for file in file_dir:
            add_file = os.path.join(path, file)
            zf.write(add_file)


def main():
    # unzip_archive()
    # d_input, d_replace = input_data()
    # replace_data(d_input, d_replace)
    # del_paragrafs_docx()
    # zip_template()
    general_format()


if __name__ == '__main__':
    main()
